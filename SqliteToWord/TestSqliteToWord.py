#! /usr/bin/env python3
# -*- coding: utf-8 -*-
import sqlite3
import os
import xml.etree.ElementTree as ET
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.shared import RGBColor
import datetime
import sys,getopt


#BIGANT5_DB = "D:/878DF07E-FA69-E229-E728-142691E5E3B0.db"
BIGANT5_DB = "D:/Documents/BigAnt5/Data/878DF07E-FA69-E229-E728-142691E5E3B0/251/878DF07E-FA69-E229-E728-142691E5E3B0.db"

#BIGANT5_DB = "D:/Documents/BigAnt5/Data/878DF07E-FA69-E229-E728-142691E5E3B0/251/878DF07E-FA69-E229-E728-142691E5E3B0 - 副本.db"

BIGANT5_FILE_SQL = "D:/test_sql.txt"

CURRENT_USER_NAME = "李青健"
CURRENT_USER_ID = "251@878DF07E-FA69-E229-E728-142691E5E3B0"


def dbToSql(con,to_file):
    """备份数据库为SQL语句"""
    # con = sqlite3.connect('d:\sqlite3_test.db')
    # 'd:\sqlite3_test_dump.sql'
    with open(to_file, 'w',encoding="utf8") as f:
        for line in con.iterdump():
            f.write('%s\n' % line)
    con.close()        

#dbToSql()
            
def getConn(db_name):
    # if os.path.isfile():
    #   os.remove(db_file)
    #db_name = "D:\878DF07E-FA69-E229-E728-142691E5E3B0.db"
    if db_name is None:
        db_name = 'd:\sqlite3_test.db'
    return sqlite3.connect(db_name)

def back_1():
    conn = getConn("D:\878DF07E-FA69-E229-E728-142691E5E3B0.db")
    dbToSql(conn,to_file="""D:/bigant5_to_sql_20190126.sql""")
    conn.close()
      
##back_1()


def cre_ins_t():
    conn = getConn()  
    cursor = conn.cursor()    
    cursor.execute('create table user (id varchar(20) primary key, name varchar(20))')
    cursor.execute('insert into user (id, name) values (\'1\', \'Michael\')')
    print("rowcount ",cursor.rowcount)
    cursor.close()
    conn.commit()
    conn.close()

# cre_ins_t();

def open_sql(conn,sql,params):
    #print("sql="+sql)
    ##print("params="+str(params))
    cursor = conn.cursor() 
    if params:
        cursor.execute(sql,params)
    else:
        cursor.execute(sql)
    return cursor;


def show_tab():
    conn = getConn() 
    #print(dir(conn))
    try:
        cursor = open_sql(conn,'select * from user where id=?', ('1',))
        #print(cursor.fetchall())
        #print(dir(cursor))
        
        for row in cursor:
            print("id=",row[0],",name=",row[1], sep='')
        
    except Exception as e:
        print('Error:',e)
    finally: 
        if cursor:
            cursor.close()
    conn.close()

##show_tab();
# dir(sqlite3.Cache)   



def dict_factory(cursor, row):
    """ 设置输出游标行 字典名称为字段名 """
    d = {}
    for idx, col in enumerate(cursor.description):
        d[col[0]] = row[idx]
    return d

def test_memory_dict_factory():
    con = sqlite3.connect(":memory:")
    con.row_factory = dict_factory
    cur = con.cursor()
    cur.execute("select 1 as a,2 as b")
    print(cur.fetchone())
    #print(cur.fetchone()["a"])
    cur.close()
    con.close()


#test_memory_dict_factory() 



def getSql(): 
    f = open(BIGANT5_FILE_SQL,"r",encoding="utf-8") 
    return "".join( f.readlines() ).split(';')


def get_sql1():
    sql_member=''
    sql_msg=''
    for i,s in enumerate(getSql()):
        if i==23:
            sql_member=s
        if i==25:
            sql_msg=s
        
        #print(i," = ",s,";",sep='')
    
    print(sql_member)
    print(sql_msg)    
    conn = getConn(db_name=BIGANT5_DB)
    conn.row_factory = dict_factory
    cs = {"rq":"2019-01-23"}
    c_member = open_sql(conn,sql_member,cs)
    for row in c_member:
        #print(row)
        params = {"SENDER_ID": row['id'], "SENDER_NAME": row['name']
                 ,"RECV_ID": CURRENT_USER_ID, "RECV_NAME":CURRENT_USER_NAME
                 ,"rq":cs["rq"]}
        c_msg_count = open_sql(conn,sql_msg,params)
        for c_msg_row in c_msg_count:
            print(row['name'],"<-->",CURRENT_USER_NAME,"count:",c_msg_row["count"])
              

    



from docx.enum.style import WD_STYLE_TYPE
 
document = Document()
styles = document.styles
 
style = document.styles['Normal']
para_format = style.paragraph_format
#print(dir(para_format),sep="\n")
para_format.space_after = Pt(1) 
para_format.space_before = Pt(1)
para_format.line_spacing = Pt(1);
para_format.line_spacing_rule = 0;
   
document.add_paragraph('Paragraph style is 中国 :aa ')

#styleT3 = None;
#生成所有段落样式
# for s in styles:
#     if s.type == WD_STYLE_TYPE.PARAGRAPH:
#         document.add_paragraph('Paragraph style is 中国 : '+ s.name, style = s)
#         if (s.name == "Heading 3"):
#             styleT3 = s;
        #print(s)
  
#document.save('d:\\para_style.docx')

def add_Text_docx(doc,mess,fontstyle=None,titlestyle=None):
    if doc == None:
        return  
    #paragraph = doc.add_paragraph()  
    #paragraph.space_after = Pt(5) 
    #paragraph.space_before = Pt(5),style='Normal'
    
    if titlestyle == None:    
        run = doc.add_paragraph('').add_run(mess)  # 字符样式
        run.font.name=u'宋体'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        if (fontstyle == "text" or fontstyle==None or fontstyle=="at"):
            run.font.bold = fontstyle == "at"
            run.font.color.rgb = RGBColor(0, 0, 0)
        elif (fontstyle in ("attach","attach_nodown") ):
            run.font.bold = fontstyle == "attach_nodown"
            run.font.color.rgb = RGBColor(192, 80, 0)
        elif (fontstyle=="user"):
            run.font.color.rgb = RGBColor(57, 108, 191)
        elif (fontstyle=="current_user"):
            run.font.color.rgb = RGBColor(0, 128, 0)
        elif (fontstyle=="title"):
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 109, 254)
    else:
        doc.add_paragraph(mess,style=titlestyle) #.add_run(mess)
        

def add_Image_docx(doc,filename):
    if doc == None:
        return
    try:
        doc.add_picture(filename)
        return True
    except Exception as e:
        print(e)
    return False



def show_tab_local(conn,sql,params,colName):
    #conn = getConn(db_name=BIGANT5_DB)
    #print(sql)
    #print(params)
    try:
        cursor = open_sql(conn,sql, params)
        out = ""
        for row in cursor:
            val = row[colName]
            # print("val=",val)
            #print(type(val))
            if isinstance(val,bytes):
                val2 = val.decode('utf-8')
            else: #if isinstance(val,str):                
                val2 = str(val)
            #print("val2="+val2)
            if out!="":
                out = out + "\n"
            out = out + val2
        
    #except Exception as e:
    #    print('Error:',e)
    finally: 
        if cursor:
            cursor.close()
    #conn.close()
    return out


def show_attach_text(msg_id, where):
    conn = getConn(db_name=BIGANT5_DB)
    conn.row_factory = dict_factory
    #msg_id='66FEE5CD-1FA1-11E9-848F-99944746267A';
    sql = "SELECT Content_data FROM ant_Msg_Content a where 1=1 " 
    if where is not None and where !="":
        sql = sql + " " + where 
    elif msg_id is not None and msg_id !="":
        sql = sql + " and msg_id='" + msg_id + "'"
    val = show_tab_local(conn,sql,{"msg_id":msg_id},"Content_data")
    print(val)
    if where is None and msg_id is not None and msg_id !="":
        tree = ET.ElementTree(ET.XML(val));
        for elem in tree.iter():
            tag = elem.tag
            arr = elem.attrib
            print(tag,arr,elem.text)
            
def getFileFullName(conn,params):
    sql = "SELECT Local_Path FROM ant_Msg_Attach a where a.MSG_ID=:msg_id and a.File_Name=:File_Name" 
    val = show_tab_local(conn,sql,params,"Local_Path")
    return val               
            
            
def getCoutentData(conn,val,params,doc):
    """ val (ant_Msg_Content.Content_data)
          无文件时：<BTF><Font name="宋体" size="15" clr="0" flags="1" /><Text>XXX</Text></BTF>
          有文件时：<BTF><Font name="宋体" size="15" clr="0" flags="1" /><File name="F1E98F5A-1F8A-11E9-8A2D-94C691204229.png" size="102865"  fshost="200.10.10.200:6667" type="0">F1E98F5A-1F8A-11E9-8A2D-94C691204229.png;66FEE5CE1FA111E9848F99944746267A</File></BTF>
    """
    tree = ET.ElementTree(ET.XML(val));
    output = ""
    for elem in tree.iter():
        tag = elem.tag;
        if not (tag in ["BTF","Font"]):
            if (tag == "Text" and elem.text):
                #output = output + str(elem.text)
                add_Text_docx(doc,str(elem.text),"text")
            elif (tag == "at"):
                at = str(elem.text);
                try:
                    at = at.split(sep=";")[1];
                except Exception:# as e:
                    pass                 
                add_Text_docx(doc,"@{}".format(at) ,"at") 
            elif (tag == "File"):
                att = elem.attrib
                fsize = round(int(att["size"])/1024,2)
                filename = getFileFullName(conn,{"msg_id":params["msg_id"],"File_Name":att["name"]})
                filestatu = "不存在"
                filestyle = "attach"
                if (filename == None or filename == ""):
                    #print({"msg_id":params["msg_id"],"File_Name":att["name"]})
                    filename = att["name"]
                    filestatu = "未下载"
                    filestyle = "attach_nodown"
                else:
                    isloadfile = False
                    ext = os.path.splitext(filename)[1]
                    if (os.path.isfile(filename)):
                        filestatu = ""
                        if (ext==None or ext=='' or ext=='.png' or ext=='.jpg'):
                            isloadfile = add_Image_docx(doc,filename)
                            if isloadfile:
                                filestatu = "显示"
                    else:
                        filestatu = "不存在"

                add_Text_docx(doc,"[({}KB) 文件:{}] {}".format(fsize, filename, filestatu) ,filestyle) 
               
                #output = output + "\n[("+str(fsize)+"KB) 文件:"+filename+"] "+filestatu+"\n"
                #if not isloadfile:
                #    add_Text_docx(doc,"[("+str(fsize)+"KB) 文件:"+filename+"] 不存在","attach")
                  
            #print(elem.tag, elem.attrib, elem.text)
        #print(output)
    return output            
            
def getHasAttachText(conn,params,doc):
    sql = "SELECT Content_data FROM ant_Msg_Content a where a.MSG_ID=:msg_id"  
    val = show_tab_local(conn,sql,params,"Content_data")
    #print(val)
    return getCoutentData(conn,val,params,doc)


###print(getHasAttachText('7F174580-F19E-11E6-8F37-34689536AB04'))


def getMessText(conn,params,doc):
    # rq,MSG_ID,Source_ID,Subject,Sender_ID,Sender_Name,Send_Date,Data_Path,Type,Ext_Type,Read_State,Attach_Count,Flag,State,Is_RECV,RECV_ID,RECV_Name,Folder_Name,Ext_Flag
    userMessSql = """
        select a.*,b.Content_data 
        FROM (select datetime(a.Send_Date/1000000, 'unixepoch', 'localtime') rq,a.*  from ant_UserMsg a) a
          left join ant_Msg_Content b on a.msg_id=b.msg_id
        where Substr(rq,1,10) = :rq
          and ((a.Sender_ID = :SENDER_ID and a.Sender_Name=:SENDER_NAME
              and a.RECV_ID = :RECV_ID and a.RECV_Name=:RECV_NAME)
            or (a.Sender_ID = :RECV_ID and a.Sender_Name=:RECV_NAME
              and a.RECV_ID = :SENDER_ID and a.RECV_Name=:SENDER_NAME ))
        order by rq
    """
    try:
        c_msg_cursor = open_sql(conn,userMessSql, params)
        c_str = ""
        for c_msg_row in c_msg_cursor:
            #c_str = c_str + "\n" + c_msg_row["Sender_Name"] + " (" + c_msg_row["rq"] +")\n";
            username = c_msg_row["Sender_Name"]
            add_Text_docx(doc,"{} ({})".format(username ,c_msg_row["rq"])
                          ,("current_user" if username==CURRENT_USER_NAME else "user")
                          )
            
            #c_str = c_str + 
            getCoutentData(conn,c_msg_row["Content_data"],{"msg_id":c_msg_row["MSG_ID"]},doc)
#             isAttach = c_msg_row["Attach_Count"]
#             if (isAttach>0):
#                 c_str = c_str + getHasAttachText(conn,{"msg_id":c_msg_row["MSG_ID"]},doc)
#             else:
#                 c_str = c_str + c_msg_row["Subject"]
#                 add_Text_docx(doc,c_msg_row["Subject"],"text")
            
        return c_str    
        
    except Exception:# as e:
        print("sql="+userMessSql)
        print("params="+str(params))
        #print('Error:',e)
        raise        
    finally: 
        if c_msg_cursor:
            c_msg_cursor.close()

def get_week_day(date):
    week_day_dict = {
      0 : '星期一',
      1 : '星期二',
      2 : '星期三',
      3 : '星期四',
      4 : '星期五',
      5 : '星期六',
      6 : '星期天',
    }
    day = date.weekday()
    return week_day_dict[day]
  
def getSessionBigAnt5(ksrq,jsrq,fileName):
#     rq = '2019-01-24'
#     sessionSql = """
#       select a.sender_id,a.Sender_Name,a.RECV_ID,a.RECV_Name,count(1) c,count(case when a.attach_count>0 then 1 end) has_attach 
#         ,MIN(rq) I,MAX(rq) X
#       FROM (select a.*,datetime(a.Send_Date/1000000, 'unixepoch', 'localtime') rq  from ant_UserMsg a) a
#       where Substr(rq,1,10) = ?
#       group by a.sender_id,a.Sender_Name,a.RECV_ID,a.RECV_Name
#       order by i
#     """
    
    # 2019-01-24
    # '251@878DF07E-FA69-E229-E728-142691E5E3B0' and Sender_Name='李青健'
    
    sessionUserSql = """
    with a as (select a.sender_id id,a.Sender_Name name
        FROM (select a.*,datetime(a.Send_Date/1000000, 'unixepoch', 'localtime') rq  from ant_UserMsg a) a
        where Substr(rq,1,10) = :rq
          and not (Sender_ID='{id}' and Sender_Name='{name}')
        group by a.sender_id,a.Sender_Name)
      ,b as (
        select a.RECV_ID id,a.RECV_Name name
        FROM (select a.*,datetime(a.Send_Date/1000000, 'unixepoch', 'localtime') rq  from ant_UserMsg a) a
        where Substr(rq,1,10) = :rq
          and not (RECV_ID='{id}' and RECV_Name='{name}')
        group by a.RECV_ID,a.RECV_Name
      )
    """.format(id=CURRENT_USER_ID, name=CURRENT_USER_NAME)
    
    resultSql = """
    select * from a 
    where exists(select 1 from b where a.id=b.id and a.name=b.name) 
    union all
    select * from b 
    where not exists(select 1 from a where a.id=b.id and a.name=b.name)
    union all
    select * from a 
    where not exists(select 1 from b where a.id=b.id and a.name=b.name)     """
       
    
    groupSql = """
    select datetime(a.Send_Date/1000000, 'unixepoch', 'localtime') rq,a.*{} from ant_GroupMsg a  
    where datetime(a.Send_Date/1000000, 'unixepoch', 'localtime') like '{}%' {}
    """
       
    """
      <BTF><Font name="宋体" size="15" clr="0" flags="1" /><at>160;张勇</at><Text>勇哥，应该是系统问题吧，之前没有结算过的</Text></BTF>
      <BTF><Font name="宋体" size="15" clr="0" flags="1" /><Text>是不是可以直接修改？</Text></BTF>
      <BTF><Font name="宋体" size="15" clr="0" flags="1" /><File name="F8A433C6-0E6A-11E9-9250-001BB9DED73E.png" size="5096"  fshost="FileServer_Default" type="2">F8A433C6-0E6A-11E9-9250-001BB9DED73E.png;096095890E6B11E99250001BB9DED73E</File><Text>这个是之前办结的，都是按自费处理</Text></BTF>
    """  
    
    if os.path.isfile(fileName):
        print("文件["+fileName+"]已存在 ，退出！")
        return 
    
    doc = Document()
    #run = doc.add_paragraph().add_run("") #
    #run.font.name=u'宋体'
    #r = run._element
    #print(dir(r.rPr.rFonts))
    #return
    
    style = doc.styles['Normal']     # 正文 样式
    para_format = style.paragraph_format  # 段落格式
    ##print(dir(para_format),sep="\n")
    #print(para_format.space_after)
    para_format.space_after = Pt(1)       # 间距-段前 x 磅
    para_format.space_before = Pt(1)      # 间距-段后 x 磅
    para_format.line_spacing = Pt(1);     # 间距-行距 设置值 x 倍 , 注意：此值一定要在 line_spacing_rule 前面设置，否则会有问题
    para_format.line_spacing_rule = 0;    # 间距-行距 0-单倍行距 ... text.WD_LINE_SPACING.SINGLE , 0-5
    #doc.add_paragraph('间距-段前 x 磅');
    #doc.add_paragraph('Paragraph style is 中国 :aa ')
    #doc.save(fileName)  
    #return
    
    conn = getConn(db_name=BIGANT5_DB)
    conn.row_factory = dict_factory
    try:
        while ksrq <= jsrq:               
            rq = ksrq.isoformat() # yyyy-mm-dd 格式     
            params = {"rq":rq}
            
            curr_row = 0
            
            # 讨论组会话数量
            sessionCount_g = int(show_tab_local(conn," select count(distinct Group_Name) sl from ("+groupSql.format("",rq,"")+") t",None,"sl"))
            
            # 用户组会话数量
            sessionCount = int(show_tab_local(conn,sessionUserSql+" select count(1) sl from ("+resultSql+") t",params,"sl"))
            
            print("====start {} {} {}".format(rq,get_week_day(ksrq),sessionCount_g+sessionCount))
            
            if (sessionCount<=0):
                print("日期 " + rq + " 无用户会话！")
            else:
                add_Text_docx(doc,"{} {} {}".format(rq,get_week_day(ksrq),sessionCount_g+sessionCount),"title","Heading 2")
                
                cursor = open_sql(conn,sessionUserSql+ " " + resultSql, params)
                try:                 
                    curr_row = 0     
                    for row in cursor:
                        curr_row = curr_row + 1;
                        print("用户会话:"+str(curr_row) +"/" + str(sessionCount) + " " + row['name'])
                        
                        params = {"SENDER_ID": row['id'], "SENDER_NAME": row['name']
                                 ,"RECV_ID": CURRENT_USER_ID, "RECV_NAME":CURRENT_USER_NAME
                                 ,"rq":rq}
                        
                        add_Text_docx(doc,"[用户] {} * {}/{}".format(row['name'],curr_row,sessionCount) + ""
                                      ,"title","Heading 3")
                        
                        #mess = 
                        getMessText(conn,params,doc)
                        #mess = "\n\n ** {}  ".format(row['name']) + "\n" + mess
                        #print(mess)
                finally: 
                    if cursor:
                        cursor.close()
            #### group message #####
            
            sessionCount = sessionCount_g # int(show_tab_local(conn," select count(distinct Group_Name) sl from ("+groupSql.format("",rq,"")+") t",None,"sl"))
            if (sessionCount<=0):
                print("日期 " + rq + " 无讨论组会话！")
            else:
                if curr_row==0:
                    add_Text_docx(doc,"{} {} {}".format(rq,get_week_day(ksrq),sessionCount),"title","Heading 2")
                  
                # ,row_number() over(partition by a.group_name) n
                cursor = open_sql(conn,groupSql.format("",rq," order by a.group_name,a.send_date"), params)
                try:
                    curr_row = group_curr_row = 0                
                    p_group_name = ""     
                    for row in cursor:
                        curr_row = curr_row + 1;
                        if (curr_row==1 or p_group_name!=row["Group_Name"]):
                            p_group_name = row["Group_Name"]
                            group_curr_row = group_curr_row  + 1;
                            print("讨论组会话:"+str(group_curr_row) +"/" + str(sessionCount) + ' ' + p_group_name)
                            add_Text_docx(doc,"[讨论组] {} * {}/{}".format(p_group_name,group_curr_row,sessionCount) + ""
                                          ,"title"
                                          ,"Heading 3")
                         
                        
                        username = row["Sender_Name"]
                        add_Text_docx(doc,"{} ({})".format(username ,row["rq"])
                                  ,("current_user" if username==CURRENT_USER_NAME else "user"))   
                        getCoutentData(conn,row["MSG_Data"],{"msg_id":row["MSG_ID"]},doc)
                        
                finally: 
                    if cursor:
                        cursor.close()
            # 日期 增加1天
            ksrq = ksrq + datetime.timedelta(days = 1)
    
        doc.save(fileName)  
        print("\n文件 ["+fileName+"] 保存成功。")  
    #except Exception as e:
    #    print('Error:',e)
    finally: 
        if conn:
            conn.close()


# def getGroupMsg(rq):
#     groupSql = """
#     select datetime(a.Send_Date/1000000, 'unixepoch', 'localtime') rq,a.* from ant_UserMsg a  
#     where datetime(a.Send_Date/1000000, 'unixepoch', 'localtime') like '{}%'
#     order by a.send_date;
#     """.format(rq)
       
#     """
#       <BTF><Font name="宋体" size="15" clr="0" flags="1" /><at>160;张勇</at><Text>勇哥，应该是系统问题吧，之前没有结算过的</Text></BTF>
#       <BTF><Font name="宋体" size="15" clr="0" flags="1" /><Text>是不是可以直接修改？</Text></BTF>
#       <BTF><Font name="宋体" size="15" clr="0" flags="1" /><File name="F8A433C6-0E6A-11E9-9250-001BB9DED73E.png" size="5096"  fshost="FileServer_Default" type="2">F8A433C6-0E6A-11E9-9250-001BB9DED73E.png;096095890E6B11E99250001BB9DED73E</File><Text>这个是之前办结的，都是按自费处理</Text></BTF>
#     """   
 
def usage():
    print("使用： " + sys.argv[0] + " [参数] ")
    print("  例子: -k 2019-01-01 -j 2019-02-03 -o aa.docx -h --help")
    print("  -h/--help 显示帮助")    
    print("  -o 输出文件名, 不指定默认 当前目录 bigant5_ksrq_jsrq.docx")
    print("  -k 开始日期 ，格式：yyyy-mm-dd , 默认今天")
    print("  -j 结束日期，格式：yyyy-mm-dd , 默认今天")    

if __name__ == '__main__':
    #show_attach_text(None,""" and msg_id in ('66FEE5CC-1FA1-11E9-848F-99944746267A',
    #'66FEE5CD-1FA1-11E9-848F-99944746267A','DFC95E50-1FA5-11E9-90E5-005056C00008','4B36FB7C-1FA6-11E9-848F-99944746267A')""")
    #show_attach_text("66FEE5CD-1FA1-11E9-848F-99944746267A",None)
    #getSessionBigAnt5("2019-03-02","d:/bigant5_2019-03-02_all.docx")
    
    ksrq = datetime.date.today()
    jsrq = datetime.date.today()
    fileName = None
    opts, args = getopt.getopt(sys.argv[1:], "ho:k:j:", ["help"])
    for op, value in opts:
        if op=="-o":
            fileName = value
        elif op=="-k":
            ksrq = datetime.datetime.strptime(value,"%Y-%m-%d").date()
        elif op=="-j":
            jsrq = datetime.datetime.strptime(value,"%Y-%m-%d").date()
        elif op == "-h" or op=="--help":
            usage()
            sys.exit()
    if ksrq>jsrq:
        print("结束日期大于开始日期，请重新操作。")
        sys.exit()
        
    if fileName is None or fileName=='':
        fileName = "bigant5_{}_{}.docx".format(ksrq.isoformat(),jsrq.isoformat())
        
    print("ksrq={},jsrq={},fileName={}".format(ksrq,jsrq,fileName))
    getSessionBigAnt5(ksrq,jsrq,fileName)




#getSessionBigAnt5("2019-01-25","d:/bigant5_2019-01-25_a1.docx")

#import datetime

# curDay = datetime.date.today();
# curDay1 = curDay + datetime.timedelta(days = 1)
# print(curDay)
# print(curDay1)

#ksrq = datetime.date(2019,1,23)
#jsrq = datetime.date(2019,1,25)
#fileName = "d:/bigant5_{}-{}.docx".format(ksrq.isoformat(),jsrq.isoformat())
#getSessionBigAnt5(ksrq,jsrq,fileName)

#d2 = datetime.datetime.strptime("2019-03-12","%Y-%m-%d");
#print(d2 , d2 + datetime.timedelta(days = 1))


#datetime.date(2017,3,22).strftime("%Y%m%d") #'20170322'
#datetime.time(12,20,59,899).strftime('%H:%M:%S') # '12:20:59'
#datetime.time(12,20,59,899).isoformat() #'12:20:59.000899'
    
    ## 通过导入 builtins 模块，可以获得内置函数、异常和其他对象的列表
#     import builtins
#     for s in dir(builtins):
#         if not s.startswith("__"):
#             pass #print(s)
        
    #print(getmembers(biltins))
    #import inspect
    #print(inspect.getmembers(1))
#     for s in inspect.getmembers(""):
#         if not s[0].startswith("__"):
#             print(s)

#     print("\n\n  start str  ")
#     d = Document()
#     print(d.__dict__)
#     for s in d.__dict__:
#         print(s)

    # print(object.__dict__.keys())
    #print(inspect._main())
    #print(help(inspect.getmembers))